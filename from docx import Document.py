import requests
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# GraphQL 查询
query = """
{
  items(lang: zh) {
    id
    name
    description
    shortName
    types
    updated
    velocity
    weight
    width
    normalizedName
    loudness
    lastOfferCount
    lastLowPrice
    avg24hPrice
    high24hPrice
  }
}
"""

# 中文字段映射
columns_zh = {
    "id": "ID",
    "name": "名称",
    "shortName": "简称",
    "description": "描述",
    "types": "类型",
    "updated": "更新时间",
    "velocity": "初速",
    "weight": "重量（kg）",
    "width": "宽度（格）",
    "normalizedName": "标准名称",
    "loudness": "响度",
    "lastOfferCount": "当前上架数量",
    "lastLowPrice": "最低价格",
    "avg24hPrice": "24小时均价",
    "high24hPrice": "24小时最高价"
}

# 请求数据
def fetch_items():
    url = "https://api.tarkov.dev/graphql"
    response = requests.post(url, json={"query": query})
    data = response.json()
    return data["data"]["items"]

# 写入字段
def write_item(doc, item):
    doc.add_heading(item.get("name", "未知物品"), level=1)
    for key, label in columns_zh.items():
        value = item.get(key)
        if isinstance(value, list):
            value = ", ".join(value)
        elif isinstance(value, float):
            value = round(value, 2)
        if value is not None:
            doc.add_paragraph(f"{label}：{value}")
    doc.add_page_break()

# 设置标题
def add_title_page(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("逃离塔科夫 物品信息库")
    run.bold = True
    run.font.size = Pt(36)
    doc.add_paragraph("数据来源：tarkov.dev API")
    doc.add_page_break()

# 主流程
def main():
    print("📡 正在获取塔科夫物品数据...")
    items = fetch_items()
    print(f"✅ 获取成功，共 {len(items)} 项")

    doc = Document()
    add_title_page(doc)

    for item in items:
        write_item(doc, item)

    output_path = "塔科夫物品信息库.docx"
    doc.save(output_path)
    print(f"📄 已生成文档：{output_path}")

if __name__ == "__main__":
    main()
